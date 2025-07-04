from docx import Document
import pytesseract
from pdf2image import convert_from_path
from io import StringIO
import tempfile
from docx2pdf import convert as docx2pdf_convert

def extract_docx_text(docx_path, ocr=None):
    """
    Extracts text from a DOCX file.
    - If ocr=False: only extract text, no OCR fallback.
    - If ocr=True: force OCR (convert to PDF and OCR).
    - If ocr=None: try text, then OCR if empty.
    """

    def _extract_text_from_docx():
        doc = Document(docx_path)
        text_lines = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_lines.append(" ".join(text.split()))

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = " ".join(cell.text.strip().split())
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    text_lines.append(" | ".join(row_cells))

        # Deduplicate while preserving order
        seen = set()
        deduped = []
        for line in text_lines:
            if line not in seen:
                deduped.append(line)
                seen.add(line)

        return "\n".join(deduped).strip()

    def _perform_ocr():
        print("🟠 OCR processing DOCX file (convert to PDF)...")
        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = f"{tmpdir}/temp.pdf"
            # Convert DOCX to PDF
            docx2pdf_convert(docx_path, pdf_path)
            # OCR the PDF
            images = convert_from_path(pdf_path)
            ocr_text = StringIO()
            for img in images:
                ocr_text.write(pytesseract.image_to_string(img))
            return ocr_text.getvalue().strip()

    if ocr is True:
        return _perform_ocr()
    elif ocr is False:
        return _extract_text_from_docx()
    else:
        extracted_text = _extract_text_from_docx()
        if extracted_text:
            return extracted_text
        else:
            return _perform_ocr()
